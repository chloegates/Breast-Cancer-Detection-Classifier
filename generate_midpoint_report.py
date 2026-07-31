#!/usr/bin/env python3
"""Generate midpoint report Word document for Random Forest breast cancer project."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PROJECT_DIR = Path(__file__).resolve().parent
RESULTS_PATH = PROJECT_DIR / "random_forest_cv5_results.json"
OUTPUT_PATH = PROJECT_DIR / "Midpoint_Report_Random_Forest_Project.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def load_results() -> dict:
    if RESULTS_PATH.exists():
        return json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    return {}


def build_report() -> Path:
    results = load_results()
    holdout = results.get("holdout_metrics", {})
    cv_acc = results.get("cv5_metrics", {}).get("accuracy", {})
    best_params = results.get("best_hyperparameters", {})
    top_features = results.get("top_features", [])

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Midpoint Progress Report")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Breast Cancer Diagnosis Prediction Using Machine Learning\n"
        "Random Forest Project Section"
    )

    doc.add_paragraph(f"Report date: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    add_heading(doc, "Project Overview")
    doc.add_paragraph(
        "Our project applies machine learning to the Wisconsin Diagnostic Breast Cancer dataset to "
        "classify tumors as benign or malignant using measurements taken from fine-needle aspirate "
        "cell samples. At the midpoint of the semester, we have established a cleaned dataset, separated "
        "training and holdout data, and built an initial Random Forest classifier using unscaled features "
        "and five-fold cross-validation. The goal of this work is to develop a reliable predictive model "
        "that can support diagnosis research while remaining transparent enough to explain in a final "
        "course report."
    )

    add_heading(doc, "Work Completed So Far")
    doc.add_paragraph(
        "The data collection and cleaning portion of the project is complete. We began with the cleaned "
        "Wisconsin dataset stored in wdbc_clean.csv, which contains 569 samples and 30 numerical features "
        "describing cell nucleus characteristics. We confirmed that the file contained no missing values, "
        "that each patient ID appeared only once, and that diagnosis labels were limited to benign and "
        "malignant cases encoded as 0 and 1."
    )
    doc.add_paragraph(
        "Preprocessing and data splitting are also complete. We divided the data into a development set "
        "of 455 samples and a final holdout set of 114 samples using an 80/20 stratified split so that "
        "the proportion of benign and malignant cases remained consistent across both files. The development "
        "data was saved as development_unscaled.csv and the holdout data as test_unscaled_FINAL_HOLDOUT.csv. "
        "We intentionally left the features unscaled because tree-based models such as Random Forest do not "
        "require normalized input in the same way that distance-based models do."
    )
    doc.add_paragraph(
        "The Random Forest model development stage is complete for the midpoint submission. We trained a "
        "RandomForestClassifier on the development data using five-fold stratified cross-validation and "
        "GridSearchCV to tune hyperparameters. The best model was saved to the models folder as "
        "random_forest_cv5_model.joblib. Model evaluation is underway. We have already computed accuracy, "
        "precision, recall, F1-score, Matthews correlation coefficient, and ROC-AUC on the held-out test "
        "set, and we exported feature importance rankings to support interpretation of the model."
    )
    doc.add_paragraph(
        "Additional algorithms remain planned rather than finished. We intend to compare the Random Forest "
        "results against logistic regression, support vector machines, and possibly XGBoost or AdaBoost. "
        "We are also continuing our literature review and reporting work, including review of published "
        "Random Forest methodology and preparation of this midpoint document. Explainability analysis using "
        "SHAP is planned but has not yet been fully implemented."
    )
    doc.add_paragraph(
        "Team member names and individual contributions should be inserted into this section before the "
        "report is submitted. At present, the work is organized by task area rather than by person."
    )

    add_heading(doc, "Random Forest Model Summary")
    if results:
        param_text = (
            f"n_estimators={best_params.get('n_estimators')}, "
            f"max_depth={best_params.get('max_depth')}, "
            f"max_features={best_params.get('max_features')}, "
            f"min_samples_split={best_params.get('min_samples_split')}, "
            f"min_samples_leaf={best_params.get('min_samples_leaf')}"
        )
        top_feature_text = ", ".join(item["feature"] for item in top_features[:5])

        doc.add_paragraph(
            f"Our current model uses {results['data']['feature_count']} unscaled input features and was "
            f"trained on {results['data']['development_samples']} development samples. The holdout file "
            f"contains {results['data']['holdout_samples']} samples that were not used during cross-validation "
            f"or hyperparameter tuning. Validation was performed with five-fold stratified cross-validation, "
            f"and the final parameter configuration was selected using GridSearchCV based on cross-validated "
            f"accuracy. The selected settings were {param_text}."
        )
        doc.add_paragraph(
            f"On the development set, the model achieved a mean cross-validated accuracy of "
            f"{cv_acc.get('mean', 0):.2%} with a standard deviation of {cv_acc.get('std', 0):.2%}. "
            f"When evaluated on the untouched holdout set, the model reached an accuracy of "
            f"{holdout.get('accuracy', 0):.2%}, precision of {holdout.get('precision', 0):.2%}, recall of "
            f"{holdout.get('recall', 0):.2%}, F1-score of {holdout.get('f1_score', 0):.2%}, Matthews "
            f"correlation coefficient of {holdout.get('mcc', 0):.3f}, and ROC-AUC of "
            f"{holdout.get('roc_auc', 0):.3f}. These results suggest that the Random Forest is performing "
            f"strongly on this dataset, although final conclusions will depend on comparison with additional "
            f"models."
        )
        doc.add_paragraph(
            f"The features that contributed most strongly to the current model included {top_feature_text}. "
            f"This pattern is consistent with the medical intuition that size- and shape-related measurements, "
            f"especially those describing the most abnormal cell characteristics, are important for "
            f"distinguishing malignant tumors from benign ones."
        )
    else:
        doc.add_paragraph(
            "Model results were not found at the time this report was generated. Run "
            "train_random_forest_model.py before regenerating this document."
        )

    add_heading(doc, "Planned Algorithms, Libraries, and Research")
    doc.add_paragraph(
        "Our modeling plan follows a structured workflow similar to published clinical machine learning "
        "research. We are using Python with pandas and NumPy for data handling, scikit-learn for model "
        "training and evaluation, and joblib to save the trained Random Forest for reuse. We may add "
        "matplotlib or seaborn later for visual review of model behavior, and we are considering SHAP for "
        "explainability if time allows."
    )
    doc.add_paragraph(
        "The immediate focus of the project is the Random Forest classifier trained on unscaled development "
        "data with five-fold cross-validation. After the midpoint, we plan to train at least one additional "
        "baseline model such as logistic regression or a support vector machine and compare the results using "
        "the same holdout set. We may also test another ensemble method such as XGBoost or AdaBoost if the "
        "remaining schedule allows. Our reporting will continue to emphasize standard classification metrics "
        "including accuracy, precision, recall, F1-score, MCC, and ROC-AUC."
    )
    doc.add_paragraph(
        "We used the Random Forest methodology described in the Scientific Reports study titled "
        "\"A robust machine learning approach to predicting remission and stratifying risk in rheumatoid "
        "arthritis patients treated with bDMARDs\" (PubMed: https://pubmed.ncbi.nlm.nih.gov/40615474/) as a "
        "reference for our approach. Although that study addresses rheumatoid arthritis rather than breast "
        "cancer, its Random Forest workflow is directly relevant to our project. The authors compared "
        "multiple machine learning models, trained them with five-fold cross-validation, tuned "
        "hyperparameters with GridSearchCV, and evaluated performance using accuracy, precision, recall, "
        "F1-score, MCC, and ROC-AUC. In that study, Random Forest performed strongly before calibration, "
        "achieving 84.42% accuracy, an F1-score of 0.842, an MCC of 0.689, and a Brier score of 0.157. "
        "They also used SHAP to identify the most influential predictors."
    )
    doc.add_paragraph(
        "We adapted that general framework to our breast cancer project by keeping a separate holdout file "
        "that is not used during cross-validation, which is similar in spirit to the external validation "
        "strategy used in the reference paper. Our current results are stronger on the Wisconsin dataset, "
        "but we are treating those numbers as preliminary until we complete comparison models and finalize "
        "our write-up."
    )

    add_heading(doc, "Roadblocks and Limitations")
    doc.add_paragraph(
        "The main limitation at this stage is time. With the midpoint deadline approaching, we do not expect "
        "to complete extensive optimization beyond the parameter tuning already performed with GridSearchCV. "
        "More advanced steps such as probability calibration, deeper hyperparameter search, or full SHAP "
        "analysis may be partially implemented or deferred to the final submission."
    )
    doc.add_paragraph(
        "We are also still in the research phase for the broader modeling plan. We have focused first on "
        "Random Forest because it fits our dataset well and aligns with the published clinical machine "
        "learning example we reviewed, but final model selection remains open. Because the project uses a "
        "single cleaned dataset rather than multiple external cohorts, we cannot yet claim the same level "
        "of generalizability described in larger clinical studies."
    )
    doc.add_paragraph(
        "Another practical concern is class imbalance. Benign cases are more common than malignant cases in "
        "both the development and holdout sets, so accuracy alone is not enough to judge the model. We are "
        "monitoring recall, MCC, and related metrics to make sure the classifier does not perform well "
        "overall while missing too many malignant cases."
    )

    add_heading(doc, "Next Steps")
    doc.add_paragraph(
        "Before the final submission, we plan to finish reviewing the Random Forest results in more detail, "
        "train and compare at least one additional classifier, and document why we chose to use unscaled "
        "features for the tree-based model. We also need to add team member names and individual "
        "contributions to the work-completed section, finalize our conclusion, and prepare the end-of-semester "
        "report. If time remains after those core tasks, we may add explainability analysis and a more formal "
        "comparison against the methodology used in the PubMed reference."
    )

    add_heading(doc, "Conclusion")
    doc.add_paragraph(
        "At the midpoint, our Random Forest project section has a reproducible preprocessing pipeline, a "
        "trained classifier using unscaled data and five-fold cross-validation, and a saved model artifact "
        "with promising holdout performance. The remaining work will focus on comparison models, final "
        "evaluation, and reporting under a tight schedule. The referenced PubMed study provides a useful "
        "template for validation, metric selection, and explainability, and we are applying that structure "
        "to our breast cancer diagnosis prediction project."
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_report()
    print(f"Created: {path}")
